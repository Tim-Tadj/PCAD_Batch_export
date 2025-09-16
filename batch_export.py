import os
import sys
import time
import re
import warnings
import PyPDF2
from pdf2image import convert_from_path
from PIL import Image
from io import BytesIO
from docx import Document
from docx.shared import Inches
from docx.enum.style import WD_STYLE_TYPE
from docxcompose.composer import Composer
from PySide6.QtWidgets import (
    QApplication, QWidget, QVBoxLayout, QHBoxLayout, 
    QLabel, QLineEdit, QPushButton, QFileDialog, 
    QProgressBar, QCheckBox, QMessageBox
)
from PySide6.QtCore import Qt, QThread, Signal
import multiprocessing
from collections import deque
from qt_material import apply_stylesheet

CWD = os.path.dirname(sys.argv[0])
CWD = os.path.normpath(CWD)
print(CWD)

# Add pdftoppm.exe and pdfinfo.exe to the path
os.environ['PATH'] = os.environ['PATH'] + ';' + CWD

warnings.simplefilter("ignore", UserWarning)

for name, value in os.environ.items():
    print("{0}: {1}".format(name, value))

def convert_dict_to_csv(data, output_file):
    with open(output_file, 'w') as file:
        file.write('Cable,Load Maximum Demand,CB Rating,Current Capacity,MAX EF impedence,EF impedence,Result\n')
        for cable, values in data.items():
            result = 'Pass' if float(values['MAX EF impedence']) > float(values['EF impedence']) else 'Fail'
            file.write(f"{cable},{values['Load Maximum Demand']},{values['CB Rating']},{values['Current Capacity']},{values['MAX EF impedence']},{values['EF impedence']},{result}\n")

def pdf_to_txt(input_path, output_path):
    basename = os.path.basename(input_path)
    basename = os.path.splitext(basename)[0]
    data = {basename: {}}
    with open(input_path, 'rb') as file:
        reader = PyPDF2.PdfReader(file)
        text = ''
        for page in range(len(reader.pages)):
            text += reader.pages[page].extract_text()
    
    # Helper: remove thousands separators
    def _clean(num_str):
        return num_str.replace(',', '')

    match = re.search(r"Load Maximum Demand\s*:\s*([\d,]+(?:\.\d+)?)", text)
    if match:
        data[basename].update({"Load Maximum Demand": _clean(match.group(1))})
        
    match = re.search(r"Rating\s*:\s*([\d,]+(?:\.\d+)?)", text)
    if match:
        data[basename].update({"CB Rating": _clean(match.group(1))})
    else:
        match = re.search(r"Rating \(In\)\s*:\s*([\d,]+(?:\.\d+)?)", text)
        if match:
            data[basename].update({"CB Rating": _clean(match.group(1))})
    
    match = re.search(r"Trip\s*:\s*([\d,]+(?:\.\d+)?)", text)
    if match:
        data[basename].update({"CB Rating": _clean(match.group(1))})
    
    match = re.search(r"Current Capacity\s*:\s*([\d,]+(?:\.\d+)?)", text)
    if match:
        data[basename].update({"Current Capacity": _clean(match.group(1))})
    
    match = re.search(r"Max\. Circuit Impedance \(max\. Zint\)\s*:\s*([\d,]+(?:\.\d+)?)", text)
    if match:
        data[basename].update({"MAX EF impedence": _clean(match.group(1))})

    match = re.search(r"Earth Fault Loop Impedance \(Zint\)\s*:\s*([\d,]+(?:\.\d+)?)", text)
    if match:
        data[basename].update({"EF impedence": _clean(match.group(1))})
    
    return data

def batch_export_process(input_files, output_files, output_folder, progress_queue):
    import pythoncom
    import pywinauto

    pythoncom.CoInitialize()
    
    try:
        pwa_app = pywinauto.application.Application()
        total_files = len(input_files)

        for i, file1 in enumerate(input_files):
            outfile1 = output_files[i]
            pwa_app.connect(title_re="PowerCad-5 - Version*")
            window = pwa_app.window(title_re="PowerCad-5 - Version*")
            window.menu_item(u'&File->&Open Project...\tCtrl+O').select()

            window = pwa_app.window(title_re="Load Project File")
            ctrl = window['Edit']
            ctrl.set_text(file1)
            ctrl = window['&Open']
            ctrl.click()

            time.sleep(0.5)
            window = pwa_app.window(title_re="PowerCad-5 - Version*")
            window.menu_item(u'&View->Local Mode').select()
            window.menu_item(u'&File->&Print...').select()

            window = pwa_app.window(title_re="PowerCad-5 Reports")
            ctrl = window['TCheckBox44']  # 43 for PCAD 5.0.80.1, 44 for PCAD 5.0.80.2 IMPORTANT: May change in future versions
            ctrl.click()

            ctrl = window['TBitBtn6']
            ctrl.click()

            window = pwa_app.window(title_re="Print")
            ctrl = window['OK']
            ctrl.click()

            window = pwa_app.window(title_re="Save Print Output As")
            ctrl = window['Toolbar4']
            ctrl.click()

            window['Edit'].set_text(os.path.normpath(output_folder + '\\' + outfile1))

            ctrl = window['&Save']
            ctrl.click()

            window = pwa_app.window(title_re="PowerCad-5 Reports")
            window["Close"].click()

            progress_queue.put(int((i + 1) / total_files * 100))

    finally:
        pythoncom.CoUninitialize()

class BatchExportThread(QThread):
    progress_update = Signal(int)
    export_complete = Signal()

    def __init__(self, input_files, output_files, output_folder):
        super().__init__()
        self.input_files = input_files
        self.output_files = output_files
        self.output_folder = output_folder

    def run(self):
        progress_queue = multiprocessing.Queue()
        process = multiprocessing.Process(target=batch_export_process, 
                                          args=(self.input_files, self.output_files, self.output_folder, progress_queue))
        process.start()

        while process.is_alive():
            while not progress_queue.empty():
                progress = progress_queue.get()
                self.progress_update.emit(progress)
            time.sleep(0.1)

        process.join()
        self.export_complete.emit()

class ConvertPDFsThread(QThread):
    progress_update = Signal(int)
    conversion_complete = Signal()

    def __init__(self, pdf_directory):
        super().__init__()
        self.pdf_directory = pdf_directory

    def run(self):
        documents = []
        
        template_doc = os.path.join(CWD, 'template.docx')
        default_doc = os.path.join(CWD, 'default.docx')

        pdf_files = [f for f in os.listdir(self.pdf_directory) if f.endswith('.pdf')]
        total_files = len(pdf_files)

        for i, pdf_file in enumerate(pdf_files):
            document = Document(template_doc)

            section = document.sections[0]
            section.page_width = Inches(8.27)
            section.page_height = Inches(11.69)

            section.left_margin = None
            section.right_margin = None
            section.top_margin = None
            section.bottom_margin = None
            
            style = document.styles['Normal'].font.name = 'Arial'
            style = document.styles['Normal'].font.bold = True

            cable_name = os.path.splitext(pdf_file)[0]
            
            for table in document.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            if "Cable Name" in paragraph.text:
                                obj_styles = document.styles
                                obj_charstyle = obj_styles.add_style('CommentsStyle', WD_STYLE_TYPE.PARAGRAPH)
                                obj_font = obj_charstyle.font
                                obj_font.bold = True
                                obj_font.name = 'Arial'
                                paragraph.text = paragraph.text.replace("Cable Name", cable_name)
                                paragraph.style = document.styles['CommentsStyle']

            with open(os.path.join(self.pdf_directory, pdf_file), 'rb') as f:
                pdf = PyPDF2.PdfReader(f)
                total_pages = len(pdf.pages)
                page_height = Inches(9.5)
                cropped_images = []

                for page_num in range(total_pages):
                    images = convert_from_path(os.path.join(self.pdf_directory, pdf_file), dpi=150, first_page=page_num+1, last_page=page_num+1)
                    image = images[0]

                    left = 170//2
                    top = 782//2
                    right = image.width - left
                    bottom = image.height - 140//2

                    image = image.crop((left, top, right, bottom))

                    bottom = image.height
                    threshold = sum(sum(pixel) for pixel in image.crop((left, bottom - 1, right, bottom)).getdata())

                    bottom -= 1
                    while bottom > 0:
                        pixel_row = image.crop((left, bottom - 1, right, bottom))
                        pixel_sum = sum(sum(pixel) for pixel in pixel_row.getdata())

                        if pixel_sum < threshold:
                            break

                        bottom -= 1
                    if bottom < 5:
                        continue
                    cropped_image = image.crop((0, 0, image.width, bottom))

                    cropped_images.append(cropped_image)

            max_image_heights = []

            combined_height = 0
            for cropped_image in cropped_images:
                combined_height += cropped_image.height

            for cropped_image in cropped_images:
                max_image_heights.append(cropped_image.height / combined_height * page_height)

            paragraph = document.add_paragraph()
            for j, cropped_image in enumerate(cropped_images):
                image_width = max_image_heights[j] * cropped_image.width / cropped_image.height * 0.85

                image_stream = BytesIO()
                cropped_image.save(image_stream, format='PNG')
                image_stream.seek(0)

                run = paragraph.add_run()
                run.add_picture(image_stream, width=image_width)

            documents.append(document)
            self.progress_update.emit(int((i + 1) / total_files * 100))

        concatenated_document = Document(default_doc)
        composer = Composer(concatenated_document)
        for document in documents:
            composer.append(document)
        composer.save(os.path.normpath(os.path.join(self.pdf_directory, 'concatenated.docx')))
        
        self.conversion_complete.emit()

class ConvertToCSVThread(QThread):
    progress_update = Signal(int)
    conversion_complete = Signal(dict)

    def __init__(self, input_dir, output_dir):
        super().__init__()
        self.input_dir = input_dir
        self.output_dir = output_dir

    def run(self):
        data = {}
        pdf_files = [f for f in os.listdir(self.input_dir) if f.endswith('.pdf')]
        total_files = len(pdf_files)

        for i, file in enumerate(pdf_files):
            input_path = os.path.normpath(os.path.join(self.input_dir, file))
            output_path = os.path.normpath(os.path.join(self.output_dir, file))
            temp = pdf_to_txt(input_path, output_path)
            data.update(temp)
            self.progress_update.emit(int((i + 1) / total_files * 100))

        self.conversion_complete.emit(data)
        
        
class FolderBrowserApp(QWidget):
    def __init__(self):
        super().__init__()
        self.setWindowTitle("POWERCAD-5 Batch Export")
        self.setGeometry(100, 100, 400, 300)  # Increased height to accommodate new label
        self.task_queue = deque()
        
        layout = QVBoxLayout()

        input_layout = QHBoxLayout()
        self.input_label = QLabel("PCAD Folder:")
        self.input_edit = QLineEdit()
        self.input_button = QPushButton("Browse")
        self.input_button.clicked.connect(self.browse_input_folder)
        input_layout.addWidget(self.input_label)
        input_layout.addWidget(self.input_edit)
        input_layout.addWidget(self.input_button)

        output_layout = QHBoxLayout()
        self.output_label = QLabel("PDF Folder:")
        self.output_edit = QLineEdit()
        self.output_button = QPushButton("Browse")
        self.output_button.clicked.connect(self.browse_output_folder)
        output_layout.addWidget(self.output_label)
        output_layout.addWidget(self.output_edit)
        output_layout.addWidget(self.output_button)
        
        checkbox_layout = QHBoxLayout()
        self.process_files_checkbox = QCheckBox("PCAD Batch Export")
        self.convert_pdfs_checkbox = QCheckBox("Convert PDFs to Word")
        self.covert_info_to_csv_checkbox = QCheckBox("Convert Info to CSV")
        checkbox_layout.addWidget(self.process_files_checkbox)
        checkbox_layout.addWidget(self.convert_pdfs_checkbox)
        checkbox_layout.addWidget(self.covert_info_to_csv_checkbox)

        self.status_label = QLabel("Ready")
        self.status_label.setAlignment(Qt.AlignCenter)
        
        self.progress_bar = QProgressBar()
        self.progress_bar.setVisible(False)

        self.process_button = QPushButton("Execute")
        self.process_button.clicked.connect(self.process_files)

        layout.addLayout(input_layout)
        layout.addLayout(output_layout)
        layout.addLayout(checkbox_layout)
        layout.addWidget(self.status_label)
        layout.addWidget(self.progress_bar)
        layout.addWidget(self.process_button)

        self.setLayout(layout)

    def browse_input_folder(self):
        folder = QFileDialog.getExistingDirectory(self, "Select Input Folder")
        if folder:
            self.input_edit.setText(os.path.normpath(folder))

    def browse_output_folder(self):
        folder = QFileDialog.getExistingDirectory(self, "Select Output Folder")
        if folder:
            self.output_edit.setText(os.path.normpath(folder))

    def process_files(self):
        input_folder = self.input_edit.text()
        output_folder = self.output_edit.text()

        if output_folder:
            if self.process_files_checkbox.isChecked():
                pdf_files = [f for f in os.listdir(output_folder) if f.lower().endswith('.pdf')]
                if pdf_files:
                    msg_box = QMessageBox()
                    msg_box.setIcon(QMessageBox.Warning)
                    msg_box.setText("PDF files were found in the output folder.")
                    msg_box.setInformativeText("Do you want to delete these files before continuing?")
                    msg_box.setStandardButtons(QMessageBox.Yes | QMessageBox.Cancel)
                    reply = msg_box.exec()

                    if reply == QMessageBox.Yes:
                        for pdf_file in pdf_files:
                            os.remove(os.path.normpath(os.path.join(output_folder, pdf_file)))
                        print("Existing PDF files deleted.")
                    elif reply == QMessageBox.Cancel:
                        print("Operation cancelled.")
                        return

            self.progress_bar.setVisible(True)
            self.progress_bar.setValue(0)
            self.process_button.setEnabled(False)

            # Clear previous task queue and add new tasks
            self.task_queue.clear()
            
            if input_folder and self.process_files_checkbox.isChecked():
                self.task_queue.append(self.start_batch_export)
            if self.convert_pdfs_checkbox.isChecked():
                self.task_queue.append(self.start_pdf_conversion)
            if self.covert_info_to_csv_checkbox.isChecked():
                self.task_queue.append(self.start_csv_conversion)

            # Start processing tasks
            self.process_next_task()

    def process_next_task(self):
        if self.task_queue:
            next_task = self.task_queue.popleft()
            next_task()
        else:
            self.finish_processing()

    def start_batch_export(self):
        self.status_label.setText("Processing PCAD Batch Export...")
        print(f"Processing files from {self.input_edit.text()} to {self.output_edit.text()}")

        input_files = [f for f in os.listdir(self.input_edit.text()) if f.endswith(".QPJ")]
        input_files = [os.path.normpath(os.path.join(self.input_edit.text(), f)) for f in input_files]
        output_files = [os.path.splitext(os.path.basename(f))[0] for f in input_files]

        print(input_files)
        print(output_files)

        self.batch_export_thread = BatchExportThread(input_files, output_files, self.output_edit.text())
        self.batch_export_thread.progress_update.connect(self.update_progress)
        self.batch_export_thread.export_complete.connect(self.process_next_task)
        self.batch_export_thread.start()

    def start_pdf_conversion(self):
        self.status_label.setText("Converting PDFs to Word...")
        print(f"Converting PDFs to Word in {self.output_edit.text()}")
        self.progress_bar.setValue(0)
        self.pdf_conversion_thread = ConvertPDFsThread(self.output_edit.text())
        self.pdf_conversion_thread.progress_update.connect(self.update_progress)
        self.pdf_conversion_thread.conversion_complete.connect(self.process_next_task)
        self.pdf_conversion_thread.start()

    def start_csv_conversion(self):
        self.status_label.setText("Converting Info to CSV...")
        print(f"Converting Info to CSV in {self.output_edit.text()}")
        self.progress_bar.setValue(0)
        self.csv_conversion_thread = ConvertToCSVThread(self.output_edit.text(), self.output_edit.text())
        self.csv_conversion_thread.progress_update.connect(self.update_progress)
        self.csv_conversion_thread.conversion_complete.connect(self.save_csv)
        self.csv_conversion_thread.start()

    def save_csv(self, data):
        output_folder = self.output_edit.text()
        convert_dict_to_csv(data, os.path.normpath(os.path.join(output_folder, '0cable_info.csv')))
        self.process_next_task()

    def finish_processing(self):
        self.progress_bar.setValue(100)
        self.process_button.setEnabled(True)
        self.status_label.setText("Complete")
        print("All processes completed!")

    def update_progress(self, value):
        self.progress_bar.setValue(value)

if __name__ == "__main__":
    multiprocessing.freeze_support()  # Necessary for multiprocessing to work with PyInstaller
    app = QApplication([])
    window = FolderBrowserApp()
    theme = os.path.normpath(os.path.join(CWD, 'color-theme.xml'))
    apply_stylesheet(app, theme=theme, invert_secondary=True)
    window.show()
    app.exec()